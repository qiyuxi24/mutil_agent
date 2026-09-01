#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""docgen.py — 文档生成工具（Markdown/HTML → DOCX/PDF）

面向 Agent 团队的确定性文档生成脚本（L3 · Skill scripts 层，见 design/TOOLCHAIN.md）：
  - Word:  Markdown/HTML → .docx   （python-docx，MIT 许可，纯 Python）
  - PDF:   Markdown/HTML → .pdf    （weasyprint 主引擎，BSD 许可；缺系统依赖时降级输出 .html）

设计原则：
  1. 解耦：本文件自包含、不依赖本项目任何内部模块，可独立复制/测试。
  2. 开源优先：仅依赖宽松许可（MIT/BSD）的开源库。
  3. 优雅降级：Markdown 渲染支持 3 级引擎（markdown → markdown_it → 内置极简渲染器），
     PDF 渲染在 weasyprint 不可用时降级写出 .html（Word 能力不受任何影响）。
  4. 确定性：同样输入产出同样文档，供 Worker 产出可审计交付物。

用法：
  python docgen.py md2docx   input.md  output.docx  [--font-family 微软雅黑] [--json]
  python docgen.py md2pdf    input.md  output.pdf   [--css style.css] [--font-family ...] [--json]
  python docgen.py html2docx input.html output.docx [--font-family 微软雅黑] [--json]
  python docgen.py html2pdf  input.html output.pdf  [--css style.css] [--font-family ...] [--json]

  input 为 `-` 时从 stdin 读取。
  --json 在 stdout 输出机器可读结果 {ok, output, format, degraded, ...}，供 Agent 解析。

依赖（requirements.txt 均已列出）：
  必装：python-docx>=1.1, markdown>=3.5        （markdown 不可用时自动用 markdown-it-py / 内置极简渲染）
  可选：weasyprint>=61（PDF 主引擎；Linux 容器需系统库，见 SKILL.md / README.md）

退出码：0 成功；1 参数/内容错误；2 依赖缺失且无法降级。
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from html.parser import HTMLParser
from pathlib import Path
from typing import Optional

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
R_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"

# ============================================================
# 一、Markdown → HTML（多引擎探测 + 极简内置降级）
# ============================================================

def _md_engine() -> str:
    """探测可用的 Markdown 渲染引擎，返回其名称。"""
    try:
        import markdown  # noqa: F401
        return "markdown"
    except ImportError:
        pass
    try:
        import markdown_it  # noqa: F401
        return "markdown_it"
    except ImportError:
        return "builtin"


def md_to_html(text: str) -> tuple[str, str]:
    """把 Markdown 文本渲染为 HTML 片段。返回 (html, engine_name)。"""
    engine = _md_engine()
    if engine == "markdown":
        import markdown as _md
        return _md.markdown(
            text,
            extensions=["tables", "fenced_code", "toc", "sane_lists"],
        ), engine
    if engine == "markdown_it":
        from markdown_it import MarkdownIt
        md = MarkdownIt("commonmark", {"html": True}).enable("table")
        return md.render(text), engine
    return _builtin_md_to_html(text), engine


_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_FENCE_RE = re.compile(r"^(`{3,}|~{3,})")
_UL_ITEM_RE = re.compile(r"^\s*[-*+]\s+(.*)$")
_OL_ITEM_RE = re.compile(r"^\s*(\d+)[.)]\s+(.*)$")
_QUOTE_RE = re.compile(r"^\s*>\s?(.*)$")
_HR_RE = re.compile(r"^\s*([-*_])\s*\1\s*\1[\s\1_]*$")
_IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)\s]+)(?:\s+\"[^\"]*\")?\)")
_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)\s]+)(?:\s+\"[^\"]*\")?\)")
_INLINE_CODE_RE = re.compile(r"`([^`]+)`")
_BOLD_RE = re.compile(r"\*\*([^*]+)\*\*|__([^_]+)__")
_ITALIC_RE = re.compile(r"\*([^*]+)\*|_([^_]+)_")


def _inline(text: str) -> str:
    """极简行内 Markdown → HTML（代码/粗体/斜体/图片/链接）。"""
    text = _IMAGE_RE.sub(r'<img alt="\1" src="\2"/>', text)
    text = _LINK_RE.sub(r'<a href="\2">\1</a>', text)
    text = _INLINE_CODE_RE.sub(r"<code>\1</code>", text)
    text = _BOLD_RE.sub(r"<strong>\1\2</strong>", text)
    text = _ITALIC_RE.sub(r"<em>\1\2</em>", text)
    return text


def _escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _builtin_md_to_html(text: str) -> str:
    """内置极简 Markdown 子集渲染器（零依赖兜底）。

    支持：标题 / 段落 / 有序无序列表 / 引用 / 分割线 / 围栏代码块 / 表格 / 行内样式。
    覆盖 90% 文档场景，保证任何环境都能产出 HTML。
    """
    out: list[str] = []
    lines = text.splitlines()
    i, n = 0, len(lines)
    in_table = False
    table_header: list[str] = []
    table_rows: list[str] = []

    def flush_table() -> None:
        nonlocal in_table, table_header, table_rows
        if not in_table:
            return
        if table_header:
            out.append("<table>")
            out.append("<thead><tr>" + "".join(f"<th>{_escape(c.strip())}</th>" for c in table_header) + "</tr></thead>")
            out.append("<tbody>")
            for row in table_rows:
                out.append("<tr>" + "".join(f"<td>{_escape(c.strip())}</td>" for c in row) + "</tr>")
            out.append("</tbody></table>")
        in_table = False
        table_header, table_rows = [], []

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # 表格行（| a | b |）
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c for c in stripped.strip("|").split("|")]
            if cells and re.match(r"^\s*:?-{3,}:?\s*$", cells[0].strip()):
                i += 1  # 分隔行：表头已记录
                continue
            if not in_table:
                table_header, table_rows, in_table = cells, [], True
            else:
                table_rows.append(cells)
            i += 1
            continue
        if in_table and not stripped:
            flush_table()
            i += 1
            continue

        if not stripped:
            flush_table()
            out.append("")
            i += 1
            continue

        m = _HEADING_RE.match(line)
        if m:
            flush_table()
            out.append(f"<h{len(m.group(1))}>{_inline(m.group(2).strip())}</h{len(m.group(1))}>")
            i += 1
            continue

        if _HR_RE.match(line):
            flush_table()
            out.append("<hr/>")
            i += 1
            continue

        m = _FENCE_RE.match(line)
        if m:
            flush_table()
            lang = line[len(m.group(1)):].strip()
            i += 1
            buf = []
            while i < n and not _FENCE_RE.match(lines[i].strip()):
                buf.append(lines[i])
                i += 1
            i += 1  # 跳过闭合围栏
            cls = f' class="language-{lang}"' if lang else ""
            out.append(f"<pre><code{cls}>" + _escape("\n".join(buf)) + "</code></pre>")
            continue

        m = _QUOTE_RE.match(line)
        if m:
            flush_table()
            buf = [m.group(1)]
            i += 1
            while i < n and _QUOTE_RE.match(lines[i]):
                buf.append(_QUOTE_RE.match(lines[i]).group(1))
                i += 1
            out.append("<blockquote>" + _inline(" ".join(x.strip() for x in buf)) + "</blockquote>")
            continue

        m = _UL_ITEM_RE.match(line)
        if m:
            flush_table()
            out.append("<ul>")
            while i < n and _UL_ITEM_RE.match(lines[i]):
                out.append("<li>" + _inline(_UL_ITEM_RE.match(lines[i]).group(1)) + "</li>")
                i += 1
            out.append("</ul>")
            continue

        m = _OL_ITEM_RE.match(line)
        if m:
            flush_table()
            out.append("<ol>")
            while i < n and _OL_ITEM_RE.match(lines[i]):
                out.append("<li>" + _inline(_OL_ITEM_RE.match(lines[i]).group(1)) + "</li>")
                i += 1
            out.append("</ol>")
            continue

        # 段落：收集连续非空行
        flush_table()
        buf = [line]
        i += 1
        while (
            i < n
            and lines[i].strip()
            and not _HEADING_RE.match(lines[i])
            and not _FENCE_RE.match(lines[i].strip())
            and not _UL_ITEM_RE.match(lines[i])
            and not _OL_ITEM_RE.match(lines[i])
            and not _QUOTE_RE.match(lines[i])
            and not _HR_RE.match(lines[i])
            and not (lines[i].strip().startswith("|") and lines[i].strip().endswith("|"))
        ):
            buf.append(lines[i])
            i += 1
        out.append("<p>" + _inline(" ".join(x.strip() for x in buf)) + "</p>")

    flush_table()
    return "\n".join(out)


# ============================================================
# 二、HTML → DOCX（python-docx，自写轻量转换器）
# ============================================================

def _require_docx():
    try:
        import docx  # noqa: F401
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("缺少 python-docx 依赖：pip install python-docx") from exc


BLOCK_TAGS = {
    "h1", "h2", "h3", "h4", "h5", "h6",
    "p", "ul", "ol", "li", "table", "thead", "tbody", "tfoot",
    "tr", "th", "td", "pre", "blockquote", "hr", "div",
}
INLINE_TAGS = {"strong", "b", "em", "i", "code", "kbd", "a", "img"}
SKIP_TAGS = {"script", "style", "head", "meta", "title"}
STACK_TAGS = BLOCK_TAGS | {"a"}


class _BlockParser(HTMLParser):
    """把 HTML 转成块树：块级标签开新节点，行内标签建子节点（供 docx/pdf 共用中间层）。"""

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.root: list[dict] = []
        self.stack: list[dict] = []
        self.text_buf: list[str] = []
        self.skip_depth = 0

    def _current(self) -> Optional[dict]:
        return self.stack[-1] if self.stack else None

    def _in_pre(self) -> bool:
        return any(n.get("type") == "pre" for n in self.stack)

    def _push(self, node: dict) -> None:
        cur = self._current()
        if cur is not None:
            cur.setdefault("children", []).append(node)
        else:
            self.root.append(node)
        if node.get("type") in STACK_TAGS:
            self.stack.append(node)

    def _flush_text(self) -> None:
        text = "".join(self.text_buf)
        self.text_buf = []
        if not text:
            return
        cur = self._current()
        if cur is not None and cur.get("type") == "pre":
            cur.setdefault("children", []).append({"type": "text", "value": text})
            return
        text = text.strip()
        if not text:
            return
        cur = self._current()
        if cur is not None:
            cur.setdefault("children", []).append({"type": "text", "value": text})
        else:
            self.root.append({"type": "p", "children": [{"type": "text", "value": text}]})

    def handle_starttag(self, tag: str, attrs: list) -> None:  # noqa: D102
        if tag in SKIP_TAGS:
            self.skip_depth += 1
            return
        if self.skip_depth:
            return
        a = dict(attrs)
        if tag in BLOCK_TAGS:
            self._flush_text()
            if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
                self._push({"type": tag})
            elif tag == "hr":
                self._push({"type": "hr"})
            elif tag == "table":
                self._push({"type": "table"})
            elif tag == "tr":
                self._push({"type": "tr"})
            elif tag in ("th", "td"):
                self._push({"type": tag, "align": a.get("align")})
            elif tag == "li":
                self._push({"type": "li"})
            elif tag == "pre":
                self._push({"type": "pre"})
            elif tag == "blockquote":
                self._push({"type": "blockquote"})
            elif tag in ("ul", "ol"):
                self._push({"type": tag})
            elif tag in ("thead", "tbody", "tfoot", "div"):
                self._push({"type": tag, "_container": True})
            elif tag == "p":
                self._push({"type": "p"})
        elif tag in INLINE_TAGS:
            if tag == "code" and self._in_pre():
                return  # pre 内 code 只是代码块的一部分，不建节点
            if tag == "img":
                self._flush_text()
                self._push({"type": "img", "src": a.get("src", ""), "alt": a.get("alt", "")})
            elif tag == "a":
                self._flush_text()
                self._push({"type": "a", "href": a.get("href", "")})
            elif tag in ("strong", "b", "em", "i", "code", "kbd"):
                self._push({"type": tag})
        elif tag == "br":
            self.text_buf.append("\n")

    def handle_endtag(self, tag: str) -> None:  # noqa: D102
        if tag in SKIP_TAGS:
            if self.skip_depth:
                self.skip_depth -= 1
            return
        if self.skip_depth:
            return
        if tag in BLOCK_TAGS:
            self._flush_text()
            while self.stack and self.stack[-1].get("type") != tag:
                if not self.stack[-1].get("_container"):
                    break  # 标签不匹配且非容器：不强行弹出，防御不规范的 HTML
                self.stack.pop()
            if self.stack and self.stack[-1].get("type") == tag:
                self.stack.pop()
        elif tag in INLINE_TAGS:
            if tag == "code" and self._in_pre():
                return
            self._flush_text()
            while self.stack and self.stack[-1].get("type") == tag:
                self.stack.pop()
        elif tag == "br":
            pass

    def handle_data(self, data: str) -> None:  # noqa: D102
        if not self.skip_depth:
            self.text_buf.append(data)


def _html_to_blocks(html: str) -> list[dict]:
    """把 HTML 片段解析为结构化块树（docx/pdf 渲染端共用的中间层）。"""
    parser = _BlockParser()
    parser.feed(html)
    parser.close()
    return parser.root


def _set_run_font(run, font_family: Optional[str], size_pt: Optional[float] = None,
                  bold: bool = False, italic: bool = False,
                  color: Optional[str] = None) -> None:
    """设置 run 字体（含中文字体 eastAsia），解决 Word 中文乱码/字体问题。"""
    run.bold = bold
    run.italic = italic
    if size_pt:
        from docx.shared import Pt
        run.font.size = Pt(size_pt)
    if color:
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor.from_string(color)
    if font_family:
        run.font.name = font_family
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(f"{W_NS}rFonts")
        if rfonts is None:
            rfonts = rpr.makeelement(f"{W_NS}rFonts", {})
            rpr.append(rfonts)
        rfonts.set(f"{W_NS}eastAsia", font_family)


def _append_run(para, text: str, font_family: Optional[str], size: float,
                bold: bool, italic: bool, color: Optional[str], code: bool) -> None:
    run = para.add_run(text)
    _set_run_font(run, font_family, size, bold=bold, italic=italic, color=color)
    if code:
        run.font.name = "Consolas"
        rpr = run._element.get_or_add_rPr()
        shd = rpr.makeelement(
            f"{W_NS}shd",
            {f"{W_NS}val": "clear", f"{W_NS}fill": "F2F2F2"},
        )
        rpr.append(shd)


def _append_hyperlink(para, url: str, text: str, font_family: Optional[str], size: Optional[float]) -> None:
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    part = para.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = para._p.makeelement(
        f"{W_NS}hyperlink",
        {f"{R_NS}id": r_id},
    )
    run = para.add_run(text)
    _set_run_font(run, font_family, size, color="0563C1")
    run.font.underline = True
    hyperlink.append(run._element)
    para._p.append(hyperlink)


def _plain_text(children: list[dict]) -> str:
    """块节点 children → 纯文本（递归）。"""
    parts = []
    for c in children or []:
        if c.get("type") == "text":
            parts.append(c.get("value", ""))
        else:
            parts.append(_plain_text(c.get("children", [])))
    return "".join(parts)


def _collect_table_rows(nodes: list[dict]) -> list[dict]:
    """从 table 子树收集所有 <tr> 节点（穿透 thead/tbody/tfoot 容器）。"""
    rows = []
    for node in nodes or []:
        if node.get("type") == "tr":
            rows.append(node)
        elif node.get("_container") or node.get("type") in ("thead", "tbody", "tfoot"):
            rows.extend(_collect_table_rows(node.get("children", [])))
    return rows


def _render_inline_children(para, children: list[dict], font_family: Optional[str],
                            base_size: float, base_bold: bool = False,
                            base_italic: bool = False, color: Optional[str] = None) -> None:
    """把块节点的 children（text + 行内节点）渲染成 docx 段落 runs。"""
    for child in children or []:
        ctype = child.get("type", "text")
        if ctype == "text":
            _append_run(para, child.get("value", ""), font_family, base_size,
                        base_bold, base_italic, color, code=False)
        elif ctype == "a":
            href = child.get("href", "")
            label = _plain_text(child.get("children", [])) or href
            _append_hyperlink(para, href, label, font_family, base_size)
        elif ctype == "img":
            src = child.get("src", "")
            if src and not src.startswith("data:"):
                try:
                    para.add_run().add_picture(Path(src))
                except Exception:  # noqa: BLE001 图片缺失不致命
                    para.add_run(f"[图片: {child.get('alt', src)}]")
        elif ctype == "br":
            para.add_run().add_break()
        elif ctype in ("strong", "b"):
            _render_inline_children(para, child.get("children", []), font_family,
                                    base_size, True, base_italic, color)
        elif ctype in ("em", "i"):
            _render_inline_children(para, child.get("children", []), font_family,
                                    base_size, base_bold, True, color)
        elif ctype == "code":
            _append_run(para, _plain_text(child.get("children", [])), font_family,
                        base_size - 1.0, base_bold, base_italic, color, code=True)
        else:
            _render_inline_children(para, child.get("children", []), font_family,
                                    base_size, base_bold, base_italic, color)


def _shade_paragraph(para, fill: str = "F5F5F5") -> None:
    """给段落加浅灰背景（代码块效果）。"""
    ppr = para._p.get_or_add_pPr()
    shd = ppr.makeelement(f"{W_NS}shd", {f"{W_NS}val": "clear", f"{W_NS}fill": fill})
    ppr.append(shd)


def _render_block_to_docx(block: dict, doc, font_family: Optional[str], default_size: float) -> None:
    """把单个块节点渲染进 docx 文档。"""
    from docx.shared import Pt

    btype = block.get("type", "p")
    children = block.get("children", [])
    text = _plain_text(children)

    if btype == "hr":
        para = doc.add_paragraph()
        run = para.add_run("─" * 40)
        _set_run_font(run, font_family, 9, color="AAAAAA")
        return

    if btype == "img":
        src = block.get("src", "")
        if src and not src.startswith("data:"):
            try:
                from docx.shared import Inches
                doc.add_picture(Path(src), width=Inches(6))
                return
            except Exception:  # noqa: BLE001
                doc.add_paragraph(f"[图片: {block.get('alt', src)}]")
                return
        return

    if btype.startswith("h") and len(btype) == 2:
        level = int(btype[1])
        para = doc.add_heading(level=min(level, 9))
        _render_inline_children(para, children, font_family, default_size + (4 - level) * 1.0)
        return

    if btype in ("ul", "ol"):
        for item in children or []:
            if item.get("type") != "li":
                _render_block_to_docx(item, doc, font_family, default_size)
                continue
            para = doc.add_paragraph(style="List Bullet" if btype == "ul" else "List Number")
            para.paragraph_format.space_after = Pt(2)
            _render_inline_children(para, item.get("children", []), font_family, default_size)
        return

    if btype == "blockquote":
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Pt(24)
        _render_inline_children(para, children, font_family, default_size, color="595959")
        return

    if btype == "pre":
        for line in (text or "").splitlines() or [""]:
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Pt(24)
            para.paragraph_format.space_after = Pt(0)
            _shade_paragraph(para)
            run = para.add_run(line)
            _set_run_font(run, "Consolas", default_size - 1.0)
        return

    if btype == "code":
        para = doc.add_paragraph()
        _shade_paragraph(para)
        run = para.add_run(text)
        _set_run_font(run, "Consolas", default_size - 1.0)
        return

    if btype == "table":
        tr_nodes = _collect_table_rows(children)
        rows = []
        for tr in tr_nodes:
            row = []
            for cell in tr.get("children", []):
                if cell.get("type") in ("th", "td"):
                    row.append(_plain_text(cell.get("children", [])))
            if row:
                rows.append(row)
        if rows:
            ncols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=ncols)
            table.style = "Table Grid"
            for ri, row in enumerate(rows):
                for ci, cell_val in enumerate(row):
                    cell = table.cell(ri, ci)
                    cell.text = ""
                    para = cell.paragraphs[0]
                    _render_inline_children(para, [{"type": "text", "value": cell_val}],
                                            font_family, default_size - 0.5)
                    if ri == 0:
                        for r in para.runs:
                            r.bold = True
        doc.add_paragraph()
        return

    # 默认 p / 其他块
    para = doc.add_paragraph()
    _render_inline_children(para, children, font_family, default_size)


def _set_docx_default_style(doc, font_family: Optional[str], default_size: float) -> None:
    """设置 docx Normal 样式默认字体（含中文字体），保证中文不乱码。"""
    from docx.shared import Pt
    style = doc.styles["Normal"]
    style.font.name = font_family or "Calibri"
    style.font.size = Pt(default_size)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(f"{W_NS}rFonts")
    if rfonts is None:
        rfonts = rpr.makeelement(f"{W_NS}rFonts", {})
        rpr.append(rfonts)
    rfonts.set(f"{W_NS}eastAsia", font_family or "SimSun")


def html_to_docx(html: str, out_path: Path, font_family: Optional[str] = None,
                 default_size: float = 10.5) -> Path:
    """HTML → .docx。返回输出文件路径。"""
    _require_docx()
    import docx

    blocks = _html_to_blocks(html)
    document = docx.Document()
    _set_docx_default_style(document, font_family, default_size)
    for block in blocks:
        try:
            _render_block_to_docx(block, document, font_family, default_size)
        except Exception as exc:  # noqa: BLE001 单块失败不致命
            document.add_paragraph(f"[渲染失败 {exc}]")
    document.save(str(out_path))
    return out_path


# ============================================================
# 三、HTML → PDF（weasyprint 主引擎 + HTML 降级）
# ============================================================

DEFAULT_CSS = """\
@page {
  size: A4;
  margin: 2cm 1.8cm;
  @bottom-center { content: counter(page) " / " counter(pages); font-size: 9pt; color: #888; }
}
body { font-family: "Noto Sans CJK SC", "Microsoft YaHei", "PingFang SC", sans-serif;
       font-size: 10.5pt; line-height: 1.6; color: #1a1a1a; }
h1 { font-size: 20pt; border-bottom: 2px solid #2c6fbb; padding-bottom: 4pt; }
h2 { font-size: 16pt; color: #2c6fbb; border-bottom: 1px solid #c9d8ea; padding-bottom: 3pt; }
h3 { font-size: 13pt; color: #2c6fbb; }
h4 { font-size: 11.5pt; }
code { background: #f2f2f2; padding: 1pt 3pt; border-radius: 2pt; font-family: Consolas, monospace; font-size: 9pt; }
pre { background: #f7f7f7; border-left: 3pt solid #2c6fbb; padding: 8pt; overflow-wrap: break-word;
      white-space: pre-wrap; font-size: 9pt; }
pre code { background: none; padding: 0; }
blockquote { border-left: 3pt solid #c9d8ea; margin: 8pt 0; padding: 4pt 10pt; color: #595959; }
table { border-collapse: collapse; width: 100%; margin: 8pt 0; font-size: 9.5pt; }
th, td { border: 1px solid #999; padding: 4pt 6pt; text-align: left; }
th { background: #eef3f9; font-weight: bold; }
tr:nth-child(even) td { background: #fafbfc; }
a { color: #0563C1; text-decoration: none; }
img { max-width: 100%; }
hr { border: none; border-top: 1px solid #bbb; margin: 10pt 0; }
"""


def _weasyprint_available() -> bool:
    try:
        import weasyprint  # noqa: F401
        return True
    except ImportError:
        return False


def _inject_font(css: str, font_family: Optional[str]) -> str:
    if not font_family:
        return css
    old = 'font-family: "Noto Sans CJK SC", "Microsoft YaHei", "PingFang SC", sans-serif;'
    new = f'font-family: "{font_family}", "Noto Sans CJK SC", "Microsoft YaHei", sans-serif;'
    return css.replace(old, new)


def html_to_pdf(html: str, out_path: Path, css: Optional[list[str]] = None,
                font_family: Optional[str] = None, allow_fallback: bool = True) -> tuple[Path, str]:
    """HTML → .pdf（weasyprint）。不可用时降级写出 .html。

    返回 (输出路径, 引擎名)。engine ∈ {"weasyprint", "html_fallback"}。
    """
    if not _weasyprint_available():
        if not allow_fallback:
            raise RuntimeError("缺少 weasyprint 依赖：pip install weasyprint")
        fallback = out_path.with_suffix(".html")
        fallback.write_text(_wrap_html(html, css, font_family), encoding="utf-8")
        return fallback, "html_fallback"

    import weasyprint

    stylesheets = [_inject_font(DEFAULT_CSS, font_family)]
    stylesheets += list(css or [])
    html_doc = weasyprint.HTML(string=_wrap_html(html, None, None), base_url=str(Path.cwd()))
    html_doc.write_pdf(str(out_path), stylesheets=[weasyprint.CSS(string=s) for s in stylesheets])
    return out_path, "weasyprint"


def _wrap_html(body: str, css: Optional[list[str]] = None, font_family: Optional[str] = None) -> str:
    """把 HTML 片段包成完整 HTML 文档（含默认 CSS）。"""
    style = _inject_font(DEFAULT_CSS, font_family)
    extra = "\n".join(f"<style>{s}</style>" for s in (css or []))
    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8"/>
<title>Document</title>
<style>{style}</style>
{extra}
</head>
<body>
{body}
</body>
</html>"""


# ============================================================
# 四、CLI
# ============================================================

def _read_input(path: str) -> str:
    if path == "-":
        return sys.stdin.read()
    return Path(path).read_text(encoding="utf-8")


def _build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        prog="docgen",
        description="文档生成工具：Markdown/HTML → DOCX/PDF（开源库，确定性输出）",
    )
    sub = p.add_subparsers(dest="command", required=True)

    for name, help_text in (
        ("md2docx", "Markdown → Word(.docx)"),
        ("html2docx", "HTML → Word(.docx)"),
    ):
        sp = sub.add_parser(name, help=help_text)
        sp.add_argument("input", help="输入文件路径（- 表示 stdin）")
        sp.add_argument("output", help="输出文件路径")
        sp.add_argument("--font-family", default=None, help="文档字体（默认微软雅黑/Noto Sans CJK）")
        sp.add_argument("--font-size", type=float, default=10.5, help="正文字号（磅，默认 10.5）")
        sp.add_argument("--json", action="store_true", help="stdout 输出机器可读 JSON 结果")

    for name, help_text in (
        ("md2pdf", "Markdown → PDF"),
        ("html2pdf", "HTML → PDF"),
    ):
        sp = sub.add_parser(name, help=help_text)
        sp.add_argument("input", help="输入文件路径（- 表示 stdin）")
        sp.add_argument("output", help="输出文件路径")
        sp.add_argument("--css", action="append", default=None, help="额外 CSS 文件（可多次指定）")
        sp.add_argument("--font-family", default=None, help="文档字体（默认 Noto Sans CJK SC/微软雅黑）")
        sp.add_argument("--no-fallback", action="store_true",
                        help="weasyprint 不可用时直接报错而非降级输出 .html")
        sp.add_argument("--json", action="store_true", help="stdout 输出机器可读 JSON 结果")
    return p


def main(argv: Optional[list[str]] = None) -> int:
    args = _build_parser().parse_args(argv)
    try:
        if args.command in ("md2docx", "md2pdf"):
            text = _read_input(args.input)
            html, engine = md_to_html(text)
            if args.command == "md2docx":
                out = html_to_docx(html, Path(args.output),
                                   font_family=args.font_family, default_size=args.font_size)
                result = {"ok": True, "output": str(out), "format": "docx",
                          "md_engine": engine, "degraded": False}
            else:
                css = [Path(c).read_text(encoding="utf-8") for c in (args.css or [])]
                out, pdf_engine = html_to_pdf(html, Path(args.output), css=css,
                                              font_family=args.font_family,
                                              allow_fallback=not args.no_fallback)
                result = {"ok": True, "output": str(out), "format": "pdf",
                          "md_engine": engine, "pdf_engine": pdf_engine,
                          "degraded": pdf_engine != "weasyprint"}
        else:  # html2docx / html2pdf
            html = _read_input(args.input)
            if args.command == "html2docx":
                out = html_to_docx(html, Path(args.output),
                                   font_family=args.font_family, default_size=args.font_size)
                result = {"ok": True, "output": str(out), "format": "docx", "degraded": False}
            else:
                css = [Path(c).read_text(encoding="utf-8") for c in (args.css or [])]
                out, pdf_engine = html_to_pdf(html, Path(args.output), css=css,
                                              font_family=args.font_family,
                                              allow_fallback=not args.no_fallback)
                result = {"ok": True, "output": str(out), "format": "pdf",
                          "pdf_engine": pdf_engine, "degraded": pdf_engine != "weasyprint"}

        if args.json:
            sys.stdout.write(json.dumps(result, ensure_ascii=False) + "\n")
        else:
            extra = "（降级：weasyprint 未安装，已输出 HTML）" if result.get("degraded") else ""
            print(f"[docgen] OK → {result['output']} ({result['format']}){extra}")
        return 0
    except RuntimeError as exc:
        print(f"[docgen] ERROR: {exc}", file=sys.stderr)
        return 2
    except Exception as exc:  # noqa: BLE001
        print(f"[docgen] ERROR: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    sys.exit(main())
